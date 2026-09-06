from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "outputs"
ASSET_DIR = OUTPUT_DIR / "report_assets"
OUTPUT_DOCX = OUTPUT_DIR / "SJT_iteration_process_metrics_report.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
BORDER = "C9D2DD"
WHITE = "FFFFFF"

QUALITY_ROWS = [
    (1, 32, 9, 0.8876, 0.5711, 0.7120, 0.7120, 0.9776),
    (2, 32, 12, 0.9246, 0.5766, 0.7301, 0.7301, 0.9739),
    (3, 32, 14, 0.9209, 0.5835, 0.7331, 0.7301, 0.9755),
    (4, 32, 17, 0.9295, 0.5847, 0.7372, 0.7301, 0.9809),
]
COST_ROWS = [
    (1, 2_963_583, 9_902_691),
    (2, 1_267_056, 4_864_428),
    (3, 1_244_164, 5_885_012),
    (4, 925_595, 2_419_154),
]


def set_run_font(run, name="Calibri", size=11, color="000000", bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent_dxa: int = 120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_paragraph_border_bottom(paragraph, color=BORDER, size="8", space="1"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def configure_styles(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("000000")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10
        style.paragraph_format.keep_with_next = True

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(3)
    run = header_p.add_run("SJT 虚拟被试迭代系统 · 技术流程报告")
    set_run_font(run, size=9, color=MUTED, bold=True)
    set_paragraph_border_bottom(header_p, color="D7DEE7", size="4", space="2")

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_before = Pt(3)
    run = footer_p.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    add_page_number(footer_p)
    run = footer_p.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def add_title_block(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("SJT 虚拟被试驱动的题目迭代流程与指标报告")
    set_run_font(run, size=23, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("当前代码逻辑、整卷评价指标与最近运行结果")
    set_run_font(run, size=13, color=MUTED)

    metadata = doc.add_table(rows=4, cols=2)
    set_table_geometry(metadata, [2700, 6660])
    set_table_borders(metadata, color="E2E7ED", size="4")
    rows = [
        ("报告日期", "2026年8月31日"),
        ("报告对象", "项目当前虚拟被试驱动的SJT题目开发与迭代流程"),
        ("主分析记录", "最近一组完整四轮整卷迭代记录：37dc1ab6…"),
        ("解释边界", "虚拟开发期指标，不替代真人样本的正式信度、效度与校标验证"),
    ]
    for row, (label, value) in zip(metadata.rows, rows):
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(label)
        set_run_font(r0, size=10, color=INK, bold=True)
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(value)
        set_run_font(r1, size=10, color="333333")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    set_paragraph_border_bottom(p, color=BLUE, size="12", space="4")
    run = p.add_run(
        "核心结论：题目层先按四项单题门槛判断合格或不合格；整卷层再用目标恢复度、构念选择性和虚拟重测ICC门槛评价候选测验。当前已删除批量淘汰功能，defer题目必须逐题处置。"
    )
    set_run_font(run, size=11, color=INK, bold=True)


def add_heading(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r0 = p.add_run(bold_prefix)
        set_run_font(r0, bold=True)
        r1 = p.add_run(text[len(bold_prefix):])
        set_run_font(r1)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_small_note(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size=9.5, color=MUTED, italic=True)
    return p


def add_callout(doc: Document, label: str, text: str, fill: str = CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120)
    set_table_borders(table, color="D6DEE8", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r0 = p.add_run(label + "：")
    set_run_font(r0, size=10.5, color=INK, bold=True)
    r1 = p.add_run(text)
    set_run_font(r1, size=10.5, color="333333")
    return table


def add_formula(doc: Document, formula: str, explanation: str | None = None):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120)
    set_table_borders(table, color="D6DEE8", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(formula)
    set_run_font(r, name="Consolas", size=11, color=INK, bold=True)
    if explanation:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(3)
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(explanation)
        set_run_font(r2, size=9.5, color=MUTED)
    return table


def write_table(doc: Document, headers: list[str], rows: Iterable[Iterable[str]], widths: list[int], *, header_fill=LIGHT_BLUE, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    header = table.rows[0]
    mark_header_row(header)
    for cell, text in zip(header.cells, headers):
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(text))
        set_run_font(r, size=font_size, color=INK, bold=True)
    for row_values in rows:
        row = table.add_row()
        for index, (cell, text) in enumerate(zip(row.cells, row_values)):
            if index % 2 == 0 and header_fill == LIGHT_BLUE:
                set_cell_shading(cell, "FBFCFD")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(text))
            set_run_font(r, size=font_size, color="333333")
    return table


def build_document():
    doc = Document()
    configure_styles(doc)
    doc.core_properties.title = "SJT虚拟被试驱动的题目迭代流程与指标报告"
    doc.core_properties.subject = "当前代码流程、指标定义与最近迭代结果"
    doc.core_properties.author = "SJT研发记录"
    doc.core_properties.comments = "虚拟开发期技术报告"
    add_title_block(doc)

    add_heading(doc, "一、当前迭代流程", 1)
    add_body(doc, "当前系统采用“题目层筛查 + 整卷层评价”的两层闭环。题目是否合格由单题指标决定；整套测验是否值得保留，则由蓝图、心理学理论约束和整卷虚拟传导指标共同决定。")
    add_callout(doc, "流程总览", "需求确认 → 蓝图与心理骨架 → 2倍候选出题 → 内容审查与冻结题库 → 虚拟被试施测 → 单题指标筛查 → 诊断与返修 → 局部复测 → 汇总重算整卷指标 → 理论约束组卷 → 平台期判断 → 输出测验")

    stages = [
        ("阶段1：需求、蓝图与题项池", "根据用户要求确定最终题量和构念结构。当前增量开发遵循候选题数量大于最终题量的原则；例如最终需要16道题时，前期先形成32道候选题，再由后续筛查和组卷选择。蓝图槽位固定目标facet、行为证据、情境机制和避免重复的约束。"),
        ("阶段2：出题、审题与冻结", "题目生成后先经过内容与构念审查，检查构念单一性、情境压力、决策张力、选项行为梯度、社会赞许风险和题目重复性。审查通过后形成冻结题库，后续虚拟施测使用同一题库版本。"),
        ("阶段3：虚拟被试统一施测", "使用固定的匹配条件施测：target、same_domain和cross_domain。主施测中每名虚拟被试对每道题回答一次；target组额外完成整卷重测，用于估计虚拟重测ICC。Neo-FFI和Mussel参照问卷属于虚拟开发期的辅助参照数据，不是正式真人校标。"),
        ("阶段4：计算单题指标并判断合格/不合格", "虚拟作答完成后，程序对全部候选题计算单题指标。四项资格门槛通过的题进入合格题集合并锁定；任意门槛未通过的题进入不合格队列。目标组选项均值梯度等指标用于定位和触发返修，不单独决定题目资格。"),
        ("阶段5：不合格题诊断与局部返修", "诊断Agent可以并发分析多道不合格题；对可返修题，Subagent执行“修改当前题 → 单题局部复测 → 不通过继续修改”的闭环，局部返修次数有上限。诊断证据不足的题进入defer，由人工逐题选择人工修改、SME审核、淘汰补题或暂停。当前已删除批量淘汰功能，因此不会再出现“一次选择影响后续所有defer题”的分支。"),
        ("阶段6：统一重算与组卷", "局部返修只用于快速试错，正式整卷指标必须由主流程统一计算。未改变的题复用已有作答数据，改变的题补充施测，然后对当前题库重新计算整卷指标。组卷Agent按蓝图题量、构念覆盖、心理学理论、机制互补和情境去重等硬约束，从候选题中选择整卷最优组合。"),
        ("阶段7：整卷迭代与平台期停止", "每一轮先形成临时测验并记录整卷指标，作为该轮候选结果；再根据单题返修结果更新题库并进入下一轮。候选整卷质量只有在超过历史最优至少0.01时才更新历史最优；连续2轮没有达到这一提升幅度时，系统判定进入平台期并停止自动返修，输出当前历史最佳组合。"),
    ]
    for title, text in stages:
        add_heading(doc, title, 2)
        add_body(doc, text)

    add_heading(doc, "二、指标体系与计算方式", 1)
    add_callout(doc, "层级关系", "单题指标回答“这道题能不能进入候选池或继续返修”；整卷指标回答“当前选出的整套测验是否比上一轮更好”。两者不能互相替代。")

    add_heading(doc, "2.1 单题资格指标", 2)
    add_body(doc, "当前单题资格由目标条件下的CITC、目标条件相关、同域VTS和跨域VTS共同决定。只要其中一项未达到门槛，该题就标记为不合格并进入诊断或返修。")
    item_rows = [
        ("Facet CITC", "Pearson(目标臂题目得分, 目标臂同facet其余题目总分)", "题目是否与同一facet的其余题目保持一致；反映局部内部一致性。", "资格门槛"),
        ("目标rho", "Spearman(Y_item,target条件分数)", "题目得分是否随目标facet设定分数有序变化；反映目标构念传导。", "资格门槛"),
        ("同域VTS", "rho_target − max(同域非目标facet的带符号rho)", "目标facet信号是否强于同一大域内其他facet信号；识别同域构念污染。", "资格门槛"),
        ("跨域VTS", "rho_target − max(跨域非目标facet的带符号rho)", "目标facet信号是否强于其他人格域信号；识别跨域构念污染。", "资格门槛"),
        ("目标组选项梯度", "目标组各选项均值是否随计分等级呈合理梯度", "定位选项是否真正表达高低行为层级；用于触发返修，不单独过滤题目。", "诊断/返修触发"),
    ]
    write_table(doc, ["指标", "计算方式", "代表的含义", "在流程中的作用"], item_rows, [1450, 3000, 3550, 1360])
    add_small_note(doc, "说明：同域和跨域VTS在每个非目标facet group内单独估计，再分别取最大带符号相关；这保证题目不是只对某一个污染facet表现良好，而是在整个匹配条件设计下具有区分性。")

    add_heading(doc, "2.2 整卷迭代指标", 2)
    add_heading(doc, "目标恢复度 R", 3)
    add_formula(doc, "R = clip(cross_validated_R2, 0, 1)", "用整套题目的完整作答模式，在留出虚拟被试上恢复预设的target构念分数。")
    add_body(doc, "计算上，系统用完整题目作答模式预测虚拟被试的target主动分数，并使用确定性的交叉验证划分。R越高，说明整套题目组合越能从作答模式恢复系统预先设定的目标构念。它是虚拟开发期的传导指标，不等同于真人样本中的汇聚效度。")

    add_heading(doc, "构念选择性 C", 3)
    add_formula(doc, "C = max(0, T) / [max(0, T) + |L|]", "T为目标构念整卷效应；L为最大非目标构念泄漏。C越接近1，目标信号占比越高。")
    add_body(doc, "其中，T定义为target组高主动分数三分位与低主动分数三分位的整卷平均分差，再除以target组整卷分数标准差；L是所有非目标条件中标准化高低组效应绝对值的最大值。C同时考虑目标信号是否存在以及非目标信号是否过强，因此比只看目标相关更能表示整卷的构念特异性。")

    add_heading(doc, "候选整卷质量 U", 3)
    add_formula(doc, "U = sqrt(R × C)", "目标恢复度与构念选择性的几何平均；两项中任意一项偏低都会拉低整卷质量。")
    add_body(doc, "U是当前整卷组卷与迭代的主要标量目标。使用几何平均的原因是避免只靠目标恢复度或只靠构念选择性取得高分：一套卷必须同时能够恢复目标构念，并且主要反映目标构念而不是邻近构念。")

    add_heading(doc, "虚拟整卷重测ICC", 3)
    add_formula(doc, "ICC = ICC(A,1) absolute agreement on repeated target form scores", "同一批target虚拟人格在同一套题目上重复施测时，总分是否稳定。")
    add_body(doc, "当前ICC只作为稳定性门槛，默认最低值为0.80，不作为U的加权组成部分。原因是虚拟被试的重复作答可能出现天花板效应，继续追求更高ICC容易把优化方向带偏。当前记录中的Cronbach α、Neo-FFI相关和Mussel相关保留为描述性诊断，不能替代真人信效度，也不作为整卷迭代目标。")

    add_heading(doc, "平台期判定", 3)
    add_formula(doc, "若 U_current ≤ U_best + 0.01，记为无显著改善；连续2轮则停止", "只有通过ICC门槛且整卷完整的轮次才进入历史最优比较。")
    add_body(doc, "因此，曲线需要同时看“本轮候选质量”和“历史最优质量”。本轮候选质量可以小幅波动或上升；历史最优质量只在超过最小改善幅度时上升，否则保持不变。")

    add_heading(doc, "三、最近四轮整卷迭代结果", 1)
    add_body(doc, "下表来自最近一组可读取的完整整卷迭代记录（运行记录37dc1ab6…）。每一轮均为16道临时测验、32道候选题池；“单题通过”表示该轮满足单题资格门槛的题目数量，不等于最终正式卷题数。")
    quality_table_rows = [
        (f"第{r[0]}轮", str(r[1]), str(r[2]), f"{r[3]:.4f}", f"{r[4]:.4f}", f"{r[5]:.4f}", f"{r[6]:.4f}", f"{r[7]:.4f}")
        for r in QUALITY_ROWS
    ]
    write_table(doc, ["轮次", "候选题", "单题通过", "目标恢复R²", "构念选择性C", "本轮质量U", "历史最优", "ICC"], quality_table_rows, [880, 1000, 1100, 1450, 1450, 1250, 1250, 980], font_size=9)
    add_small_note(doc, "解读：第1轮到第2轮U提升0.0181，超过0.01，因此历史最优从0.7120更新到0.7301；第3轮和第4轮虽然本轮U继续上升，但各自相对历史最优的增量不足0.01，所以历史最优保持0.7301。")

    add_heading(doc, "3.1 迭代成本", 2)
    cost_table_rows = [
        (f"第{r[0]}轮", f"{r[1]:,}", f"{r[2]:,}", f"{r[2] / 60000:.1f}")
        for r in COST_ROWS
    ]
    write_table(doc, ["轮次", "Token消耗", "模型耗时(ms)", "模型耗时(分钟)"], cost_table_rows, [1500, 2500, 2700, 2660], font_size=9.5)
    add_body(doc, "四轮的单轮Token消耗总体下降，主要是因为后续轮次复用了未改变题目的作答数据，并且需要处理的返修范围逐步变化。但成本不能单独证明质量提高，仍需与U、R、C及ICC门槛一起解释。")

    add_heading(doc, "3.2 最新一次新任务的基线记录", 2)
    add_body(doc, "最新检查点d34347dd…只完成了第1轮临时组卷基线，随后进入旧返修路由异常，不应与上面的四轮连续曲线直接拼接。它可以作为当前新任务的起点记录：")
    baseline_rows = [
        ("d343… 第1轮", "32", "5", "16", "0.8101", "0.5598", "0.6734", "0.9700", "5,786,351", "338.4"),
    ]
    write_table(doc, ["记录", "候选题", "单题通过", "临时测验", "目标恢复R²", "C", "U", "ICC", "Token", "耗时(分钟)"], baseline_rows, [1150, 650, 750, 750, 1000, 700, 700, 700, 1900, 1060], header_fill=LIGHT_GRAY, font_size=8.5)
    add_callout(doc, "结果边界", "这条最新基线不是质量突然下降的证据；它属于另一轮新任务，而且是在批量defer淘汰路由修复前生成的。修复后的新运行需要重新形成连续曲线，才能与37dc记录做严格的同任务比较。", fill="FFF8E8")

    add_heading(doc, "四、当前结果的综合判断", 1)
    add_body(doc, "从最近四轮完整记录看，整卷候选质量U呈数值上升，目标恢复度R²总体较高，构念选择性C连续上升，四轮ICC均通过0.80稳定性门槛。这说明在虚拟开发环境中，返修和组卷过程确实产生了改善信号。")
    add_body(doc, "但需要区分“数值递增”和“达到历史最优更新阈值”：第3、4轮的U虽然比上一轮高，但改善幅度不够0.01，因此系统仍把第2轮作为历史最优。这种设计用于避免把随机波动误判为有效进步。")
    add_body(doc, "最终问卷是否在真人上具有良好信度、汇聚效度和区分效度，不能由虚拟曲线单独证明。后续真人预测试应检验：虚拟整卷U、R²、C与真人整卷信效度之间是否具有稳定相关；这属于系统有效性验证，而不是当前虚拟迭代指标本身。")

    add_heading(doc, "五、当前实现中的关键约束", 1)
    constraint_rows = [
        ("题目资格", "单题四项门槛决定合格/不合格；合格题资格锁定，后续监测警告不撤销资格。"),
        ("返修闭环", "Subagent可并发执行单题修改—单题复测；正式整卷指标由主流程统一施测和计算。"),
        ("组卷原则", "先满足蓝图和心理学理论约束，再在可行组合中选择U较高的整卷；不按单个题目指标简单排序。"),
        ("defer处置", "每一道defer必须单独人工处理；批量淘汰开关和对应界面选项已删除。"),
        ("平台期", "连续2轮候选U没有超过历史最优至少0.01时停止自动返修并收卷。"),
        ("证据边界", "虚拟被试指标服务于系统迭代，不等于真人样本的正式心理测量结论。"),
    ]
    write_table(doc, ["约束对象", "当前规则"], constraint_rows, [1900, 7460], font_size=9.5)

    add_heading(doc, "附：主要实现位置", 1)
    add_body(doc, "整卷指标和平台期：sjt_system/evaluation/form_metrics.py；单题匹配条件指标：sjt_system/evaluation/psychometrics.py；诊断、返修与队列：sjt_system/workflow/executor.py；人工defer处置：sjt_system/workflow/interaction_nodes.py；理论组卷提示词：sjt_system/prompt/form_optimizer_prompt.py。")
    add_small_note(doc, "本报告中的数值均来自项目本地运行检查点；未对虚拟数据进行真人外推。")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_document()
