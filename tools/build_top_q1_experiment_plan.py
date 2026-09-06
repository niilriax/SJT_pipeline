from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"E:\DR_projects\langgraph_for_SJT")
OUT = ROOT / "docs" / "Top_Q1_SJT_实验方案_目标期刊与执行计划_2026-09.docx"
TABLE_WIDTH = 9360

NAVY = "0B2545"
BLUE = "2E74B5"
DEEP_BLUE = "1F4D78"
MUTED = "64748B"
PALE_BLUE = "E8EEF5"
PALE_GRAY = "F5F7FA"
PALE_GOLD = "FFF6DE"
GOLD = "B7791F"
GRID = "CBD5E1"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = "w:%s" % edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key in ["val", "sz", "space", "color"]:
            if key in kwargs[edge]:
                element.set(qn("w:%s" % key), str(kwargs[edge][key]))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:%s" % m))
        if node is None:
            node = OxmlElement("w:%s" % m)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: Sequence[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.15) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_run_font(run, size=10.5, bold=False, color="1F2937", italic=False, name="Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def add_text(paragraph, text: str, *, size=10.5, bold=False, color="1F2937", italic=False) -> None:
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)


def add_hyperlink(paragraph, text: str, url: str, *, size=9.5) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rpr.append(sz)
    run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    set_run_font(run, size=8.5, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def add_bullet(doc, text: str, *, level=0, size=10.2, color="1F2937"):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.22 * level)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    set_paragraph_spacing(p, after=3, line=1.12)
    add_text(p, text, size=size, color=color)
    return p


def add_number(doc, text: str, *, size=10.2):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    set_paragraph_spacing(p, after=3, line=1.12)
    add_text(p, text, size=size)
    return p


def add_body(doc, text: str, *, after=6, size=10.5, color="1F2937", italic=False, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=after, line=1.18)
    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, size=size, bold=True, color=DEEP_BLUE, italic=italic)
        add_text(p, text[len(bold_prefix):], size=size, color=color, italic=italic)
    else:
        add_text(p, text, size=size, color=color, italic=italic)
    return p


def add_heading(doc, text: str, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.keep_with_next = True
    if level == 1:
        set_paragraph_spacing(p, before=15, after=7, line=1.0)
        add_text(p, text, size=16, bold=True, color=BLUE)
    elif level == 2:
        set_paragraph_spacing(p, before=10, after=5, line=1.0)
        add_text(p, text, size=13, bold=True, color=DEEP_BLUE)
    else:
        set_paragraph_spacing(p, before=7, after=4, line=1.0)
        add_text(p, text, size=11.3, bold=True, color=DEEP_BLUE)
    return p


def add_callout(doc, title: str, text: str, *, fill=PALE_BLUE, title_color=DEEP_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [TABLE_WIDTH])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell,
                    top={"val": "single", "sz": "8", "color": "B8CBE0"},
                    bottom={"val": "single", "sz": "8", "color": "B8CBE0"},
                    left={"val": "single", "sz": "18", "color": BLUE},
                    right={"val": "single", "sz": "8", "color": "B8CBE0"})
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=2, line=1.12)
    add_text(p, title, size=10.5, bold=True, color=title_color)
    p2 = cell.add_paragraph()
    set_paragraph_spacing(p2, after=0, line=1.12)
    add_text(p2, text, size=10.0)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_table(doc, headers: Sequence[str], rows: Iterable[Sequence[str]], widths: Sequence[int], *, font_size=9.0, header_fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_widths(table, widths)
    table.style = "Table Grid"
    hdr = table.rows[0]
    repeat_header(hdr)
    for idx, value in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, header_fill)
        set_cell_border(cell,
                        top={"val": "single", "sz": "6", "color": GRID},
                        bottom={"val": "single", "sz": "8", "color": BLUE},
                        left={"val": "single", "sz": "4", "color": GRID},
                        right={"val": "single", "sz": "4", "color": GRID})
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, after=0, line=1.05)
        add_text(p, value, size=font_size, bold=True, color=NAVY)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            if ridx % 2 == 1:
                set_cell_shading(cell, PALE_GRAY)
            set_cell_border(cell,
                            top={"val": "single", "sz": "4", "color": GRID},
                            bottom={"val": "single", "sz": "4", "color": GRID},
                            left={"val": "single", "sz": "4", "color": GRID},
                            right={"val": "single", "sz": "4", "color": GRID})
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.08)
            add_text(p, str(value), size=font_size, color="1F2937")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_source(doc, label: str, url: str, note: str):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=3, line=1.08)
    add_text(p, f"{label}：", size=9.2, bold=True, color=DEEP_BLUE)
    add_hyperlink(p, url, url, size=8.8)
    add_text(p, f"（{note}）", size=9.0, color=MUTED)


def set_doc_defaults(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("1F2937")
    for name, size, color in (("Heading 1", 16, BLUE), ("Heading 2", 13, DEEP_BLUE), ("Heading 3", 11.3, DEEP_BLUE)):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
    for section in doc.sections:
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.72)
        section.right_margin = Inches(0.72)
        section.header_distance = Inches(0.28)
        section.footer_distance = Inches(0.28)


def add_header_footer(section, *, cover=False):
    header = section.header
    p = header.paragraphs[0]
    clear_paragraph(p)
    if not cover:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_spacing(p, after=0, line=1.0)
        add_text(p, "LLM—虚拟被试闭环 SJT 开发  |  实验方案与执行计划", size=8.2, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    clear_paragraph(fp)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(fp, after=0, line=1.0)
    add_text(fp, "研究方案确认稿  ·  2026-09-03  ·  ", size=8.2, color=MUTED)
    add_page_field(fp)


def build() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_doc_defaults(doc)
    add_header_footer(doc.sections[0], cover=True)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(9)
    add_text(p, "研究方案｜Top 一区目标", size=11, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    add_text(p, "LLM 驱动的 SJT 自动开发系统\n实验设计、期刊策略与执行计划", size=25, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    add_text(p, "验证虚拟被试闭环是否能够带来可迁移、可复现且值得其成本的测验质量提升", size=12.5, color=MUTED)
    add_callout(
        doc,
        "本方案的核心判断",
        "正式研究不应只证明系统能在虚拟被试上得到高分，而应回答：结构化流程是否优于一次性提示词；虚拟被试闭环是否产生增量价值；这种虚拟指标上的改进能否迁移到真人数据；质量提升是否值得额外的时间与 token 成本。",
        fill=PALE_BLUE,
    )
    meta = add_table(
        doc,
        ["节点", "目标", "硬性产出"],
        [
            ("2026-09-15", "完成两项实验", "系统对照结果、真人施测数据锁定或达到预注册样本量、分析脚本与原始日志归档"),
            ("2026-09-30", "完成初稿", "方法、结果、图表、限制与投稿策略形成可审阅初稿"),
        ],
        [1500, 2600, 5260],
        font_size=9.4,
    )
    for row in meta.rows:
        row.cells[0].paragraphs[0].runs[0].font.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    add_text(p, "方案状态：确认稿（结果栏位在实验完成后回填）", size=9.5, color=MUTED, italic=True)
    doc.add_page_break()

    add_heading(doc, "一、研究定位与核心结论", 1)
    add_body(doc, "研究目标不是把虚拟被试指标当作真人效度的替代品，而是检验一种“生成—施测—诊断—改题—重组”的测验开发方法是否能够提高最终真人测量质量，并降低人工试错成本。")
    add_callout(doc, "两项正式实验", "实验一检验系统内部的增量价值、稳定性与效率；实验二用真人样本检验生成问卷的真实心理测量表现。实验一是过程证据，实验二是最终有效性证据。", fill=PALE_GOLD, title_color=GOLD)
    add_heading(doc, "1.1 研究边界", 2)
    for text in [
        "Mussel 是人工开发的既有问卷，可作为真人研究中的外部参照/效标，但不是与本系统在同一蓝图、同一时间预算下开发出的“人工开发控制组”。",
        "如果论文声称“比人开发更高效”，必须增加同一蓝图下的人工专家开发条件，并记录专家工时、轮次与直接成本；否则只能声称系统优于 AI 基线，或与人工开发的 Mussel 参照进行比较。",
        "虚拟被试指标用于闭环中的筛选与诊断，真人信效度用于判断这种闭环是否具有外部效度；两者的关系必须作为可检验假设，而不能事后默认等价。",
        "若样本来自未成年人，真人施测前必须完成伦理审批、监护人同意与数据保护方案；伦理或招募未就绪时，不应把预期结果写成已完成实验。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "二、总体研究问题与预期假设", 1)
    add_body(doc, "建议把系统有效性拆成“流程增量”和“现实迁移”两层。这样可以区分：系统是否真的比简单提示词更好，以及虚拟闭环带来的改进是否在真人上仍然成立。")
    add_table(
        doc,
        ["研究问题", "要比较什么", "主要证据", "预期模式"],
        [
            ("RQ1：流程增量", "一次性提示词 vs 结构化但无虚拟闭环", "最终问卷质量、内容质量、失败率", "结构化流程优于一次性生成"),
            ("RQ2：虚拟闭环增量", "无虚拟闭环 vs 完整虚拟被试闭环", "虚拟目标传导度、构念特异度、候选质量与开发成本", "完整闭环进一步改善质量，但增加成本"),
            ("RQ3：真人迁移", "三种 AI 条件生成的问卷在真人样本上的表现", "真人信度、汇聚效度、区分效度、结构效度", "虚拟指标较高的版本在真人上也更好"),
            ("RQ4：成本—质量", "每种条件的质量增益与资源消耗", "token、时间、人工干预、轮次、质量—成本曲线", "完整闭环形成可解释的质量—成本权衡"),
        ],
        [1300, 2700, 3000, 2360],
        font_size=8.8,
    )
    add_heading(doc, "2.1 预注册的主要假设", 2)
    for text in [
        "H1：完整系统相对于一次性提示词基线，在多次独立运行和多个构念/蓝图上产生更高的真人测验质量。",
        "H2：完整系统相对于不含虚拟被试闭环的结构化流程，在虚拟目标传导度和构念特异度上更高，并在真人汇聚效度、区分效度上不劣或更优。",
        "H3：虚拟指标与真人测验指标具有正向关系；该关系须在独立运行、独立蓝图或确认样本中检验，不能只在同一批数据上回顾性拟合。",
        "H4：完整系统的质量提升伴随 token、时间或人工干预增加；研究应报告 Pareto 关系，而不是预先把多个指标任意压成一个总分。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "三、实验一：系统消融、重复性与开发效率", 1)
    add_heading(doc, "3.1 实验目的", 2)
    add_body(doc, "在完全相同的蓝图、目标构念、初始证据和最终题数要求下，比较不同自动开发流程。实验一回答“完整闭环是否有增量价值”，但不能单独证明问卷对真人有效。")
    add_heading(doc, "3.2 条件设置", 2)
    add_table(
        doc,
        ["条件", "流程", "条件中包含的机制", "能够支持的解释"],
        [
            ("A：一次性提示词", "蓝图/证据 → 一次生成 → 选出最终题目", "无结构化返修、无虚拟闭环", "最简 AI 基线"),
            ("B：结构化无闭环", "蓝图/证据 → 结构化生成与内容规则 → 组卷", "有理论与格式约束，但不使用虚拟施测驱动改题", "结构化提示与工程流程的增量"),
            ("C：完整系统", "增量出题 → 审题 → 虚拟施测 → 单题诊断/并发返修 → 统一施测 → 组卷 → 平台期停止", "完整虚拟被试闭环、固定匹配条件、质量—成本记录", "虚拟闭环相对于 B 的增量价值"),
            ("D：人工同任务对照（建议）", "专家在同一蓝图下开发同题数问卷，并记录时间和人工投入", "真人专家、同一约束、同一目标", "只有加入 D，才能严谨声称与人工开发相比的效率"),
        ],
        [1300, 2950, 2700, 2410],
        font_size=8.5,
    )
    add_callout(doc, "最低可行设计", "若当前时间无法完成 D，则保留 A/B/C 三个 AI 条件，并把 Mussel 定义为“人工开发的外部参照问卷”，不要把它写成同任务人类开发控制组。", fill=PALE_GOLD, title_color=GOLD)
    add_heading(doc, "3.3 设计控制", 2)
    for text in [
        "至少使用 3 个不同构念或蓝图；每个 AI 条件每个蓝图至少 3 次独立运行，目标为 5 次。若只能做一个构念，应增加到至少 5 个随机种子，并在论文中收窄外推结论。",
        "固定模型版本、提示词版本、温度、随机种子策略、输入证据、最终题数（例如 16 道）和输出格式；每次运行只改变实验条件允许改变的机制。",
        "同一批虚拟被试、同一组人格目标分数、同一批匹配条件和相同施测规则用于不同条件的可比指标计算。",
        "记录每轮完整的 token、时间、费用、题目数、返修轮次、人工干预、失败与重试；禁止只保留最终成功输出。",
        "内容质量可增加盲审：构念纯度、重复性、情境多样性、蓝图覆盖和选项合理性；至少两名审题者，报告一致性。",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "3.4 运行顺序", 2)
    for text in [
        "在同一蓝图下分别运行 A、B、C（如保留 D，则同步完成人工开发）。",
        "C 条件先完成初始题库和第一次虚拟施测，形成临时组卷基线；记录基线质量与成本。",
        "只对未通过或待返修题启动单题返修 subagent；每个 subagent 可完成“改题—单题局部复测—继续修改”的局部闭环。",
        "所有题目返修结束后，由主流程统一对新的完整候选题库施测并组卷；不能把不同 subagent 的局部施测直接拼成整卷数据。",
        "当连续预注册轮次数的整卷质量没有达到最小实际改善，或达到最大轮次/预算时停止；输出最终问卷与完整轨迹。",
    ]:
        add_number(doc, text)
    add_heading(doc, "3.5 实验一的结果变量", 2)
    add_table(
        doc,
        ["维度", "指标", "含义与报告方式"],
        [
            ("虚拟目标传导", "T = cor(设定目标构念分数，测验目标得分)", "目标人格水平变化是否传导为对应测验得分变化；按题目和整卷分别报告。"),
            ("虚拟构念特异度", "S = r_target − max(|r_same-domain|, |r_cross-domain|)", "目标构念的信号是否超过同域/跨域污染；越高表示选择性更强。"),
            ("虚拟候选质量", "预注册的 T、S 组合或报告向量", "用于系统内部筛选与迭代；不得未经验证直接称为真人效度。"),
            ("开发成本", "token、时间、费用、人工干预、迭代轮数", "分别报告原始值和单位合格题成本；不任意与效度合成一个分数。"),
            ("内容质量", "盲审评分与一致性", "检验重复、构念单一、情境覆盖和理论一致性。"),
        ],
        [1700, 3350, 4310],
        font_size=8.8,
    )

    add_heading(doc, "四、实验二：真人心理测量验证", 1)
    add_heading(doc, "4.1 实验目的", 2)
    add_body(doc, "实验二是论文最关键的有效性证据：让真人完成 A/B/C（以及 D，如有）的问卷，并检验虚拟被试闭环筛出的问卷是否在真实受试者上表现更好。")
    add_heading(doc, "4.2 被试与程序", 2)
    for text in [
        "在预注册中先确定主要比较、最小实际有意义差异、显著性水平、检验效能、缺失处理和剔除标准，再确定样本量；不要根据结果倒推样本量。",
        "受试者随机分配到 A/B/C（和 D）问卷版本；随机化、呈现顺序、作答环境、说明语和完成时限保持一致。",
        "所有受试者完成目标 SJT 以及 NEO-FFI、Mussel 参照问卷；如果存在长问卷负担，预先固定模块顺序或进行平衡。",
        "真人数据分析人员尽可能在条件标签盲态下运行预先写好的脚本；主结果完成后再做探索性分析。",
        "若需要验证虚拟—真人转移关系，优先使用独立确认样本；若只能拆分一个样本，明确区分探索样本与确认样本。",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "4.3 仪器与计分角色", 2)
    add_table(
        doc,
        ["问卷", "计分规则", "在实验中的角色"],
        [
            ("系统生成 SJT", "沿用蓝图中预先固定的选项分值和构念方向", "比较 A/B/C/D 条件的目标测量质量"),
            ("NEO-FFI", "按正式量表手册与预注册规则计分", "作为人格构念的参照维度，检验汇聚与区分关系"),
            ("Mussel", "0/1 二分计分；必须在代码和文档中固定该规则", "作为外部参照/效标，用于迭代系统和真人验证；不宣称为同任务人工控制组"),
        ],
        [1800, 3300, 4260],
        font_size=9.0,
    )
    add_heading(doc, "4.4 真人主要结果指标", 2)
    add_table(
        doc,
        ["结果", "计算建议", "解释"],
        [
            ("信度", "ω 为主，α 为补充；按目标构念/分量表报告置信区间", "题目是否具有足够的一致性；不是越高越好，过高可能提示题目同质化。"),
            ("汇聚效度", "系统 SJT 得分与 NEO-FFI/Mussel 对应维度的预注册相关", "是否与理论上测量相近的外部构念相关；同时报告方向和区间。"),
            ("区分效度", "目标相关与非目标相关的差值或相关模式；必要时用相关差异检验", "目标关联是否显著强于非目标关联；不能只看一个绝对相关。"),
            ("结构效度（建议纳入）", "按题数和样本允许程度，报告 CFA/EFA 或多特质—多方法证据", "检验题目结构是否支持蓝图构念；Top 一区目标下不建议完全省略。"),
        ],
        [1750, 4100, 3510],
        font_size=8.7,
    )
    add_callout(doc, "统计解释边界", "实验二的主要结论应围绕真人信度、汇聚效度和区分效度。虚拟目标传导度与构念特异度是过程指标；只有当它们能够在独立运行/样本中预测真人结果时，才可作为系统有效的先验筛选指标。", fill=PALE_BLUE)
    add_heading(doc, "4.5 建议分析", 2)
    for text in [
        "以问卷条件为固定效应、蓝图/运行/受试者为适当随机因素，或在样本量允许时使用分层模型；至少报告每个蓝图和每个条件的效应量及置信区间。",
        "比较 A→B 的流程增量和 B→C 的虚拟闭环增量，不只比较 C 与 A；这两个对比对应不同的科学问题。",
        "检验虚拟指标对真人结果的预测关系时，使用独立运行、留出蓝图或确认样本；报告相关、回归系数和预测区间。",
        "把统计显著性、实际效应、重复性和成本一起解释；不要因为信度很高就直接宣称整卷有效。",
    ]:
        add_bullet(doc, text)
    add_callout(doc, "真人实验硬门槛", "在伦理审批、知情同意、招募与数据安全未完成前，不能把真人施测写成已完成结果。若 9 月 15 日前无法完成真人数据锁定，应将该节点定义为“完成实验执行与数据收集/分析冻结”，并把未完成部分明确标为后续工作。", fill=PALE_GOLD, title_color=GOLD)

    add_heading(doc, "五、Mussel 与人工开发对照的边界", 1)
    add_body(doc, "Mussel 的作用需要写得非常准确。它可以作为既有的、人工开发的外部参照问卷，帮助检验系统生成问卷与相关人格维度之间的关系，也可以帮助系统在开发阶段进行内部迭代。但它并不能单独回答“本系统是否比人工开发更高效”。")
    add_table(
        doc,
        ["说法", "是否可支持", "原因"],
        [
            ("与人工开发的既有问卷进行外部效度比较", "可以", "Mussel 提供一个稳定参照，但需说明构念、题型、样本和计分差异。"),
            ("证明系统比人开发更快/更省", "不能只靠 Mussel", "缺少同一任务、同一蓝图、同一质量目标下的人类开发工时与成本。"),
            ("用 Mussel 0/1 分数参与相关分析", "可以，但须固定", "在代码、预注册、数据字典和论文方法中明确二分计分，避免与连续分数混用。"),
            ("用 Mussel 结果作为系统闭环的唯一优化目标", "不建议", "会把系统变成追逐单一效标，忽略目标构念纯度、区分效度和内容理论。"),
        ],
        [3000, 1800, 4560],
        font_size=8.9,
    )

    add_heading(doc, "六、目标期刊与投稿策略", 1)
    add_body(doc, "以下是按“方法学匹配度—Top 一区挑战度—现实备选”排列的投稿梯度。分区会随年份、学科类别和学校采用的榜单而变化，投稿前必须用本单位当年的 JCR/中科院分区表复核。")
    add_table(
        doc,
        ["层级", "期刊", "匹配点", "要达到的证据强度/主要风险"],
        [
            ("首选现实目标", "Behavior Research Methods", "心理学方法、计算机技术、测量工具与实验流程高度匹配", "需要多构念、多次独立运行、真人验证和透明成本报告；不要只有单个系统案例。"),
            ("Top 目标/冲刺", "Psychological Methods", "强调心理数据收集、分析、解释的方法创新，方法学要求高", "必须提出可推广的方法论贡献，并用多蓝图、多数据源和真人结果证明，不足以只展示工程实现。"),
            ("方法学备选", "Applied Psychological Measurement", "测量技术、信效度、比较验证与测验开发均有空间", "突出测量方法与验证设计，而不是把论文写成纯软件介绍。"),
            ("测量学备选", "Educational and Psychological Measurement", "覆盖心理测量理论、测验开发和创新应用", "需要更强的结构效度、测量模型或跨样本验证。"),
            ("条件性备选", "Psychological Assessment", "心理评估工具开发与验证方向相关", "临床/评估场景匹配度要求更高；若研究没有明确评估应用，适配性较弱。"),
        ],
        [1500, 2450, 2900, 2510],
        font_size=8.45,
    )
    add_heading(doc, "6.1 官方页面（投稿前复核）", 2)
    add_source(doc, "Behavior Research Methods", "https://www.psychonomic.org/page/BRM", "期刊范围与方法学定位")
    add_source(doc, "Psychological Methods", "https://www.apa.org/pubs/journals/met/", "期刊范围与方法学定位")
    add_source(doc, "Applied Psychological Measurement", "https://journals.sagepub.com/author-instructions/apm", "投稿范围与测量方法定位")
    add_source(doc, "Educational and Psychological Measurement", "https://journals.sagepub.com/home/epm", "期刊范围与测量学定位")
    add_source(doc, "Psychological Assessment", "https://www.apa.org/pubs/journals/pas", "期刊范围与评估定位")
    add_callout(doc, "投稿决策", "建议把 Behavior Research Methods 作为首选现实目标，把 Psychological Methods 作为证据充分时的冲刺目标。若没有人工同任务对照，就删去“优于人工开发效率”的措辞；若只有一个构念，也不建议将结论写成普遍适用于心理测验开发。", fill=PALE_GOLD, title_color=GOLD)

    add_heading(doc, "七、验收标准与论文可投门槛", 1)
    add_table(
        doc,
        ["层面", "最低验收标准", "Top 一区增强项"],
        [
            ("系统实验", "A/B/C 条件可复现；固定 prompt、模型、种子和输入；日志完整", "≥3 蓝图、每条件≥3 次独立运行，目标每条件 5 次；代码和提示词可审计"),
            ("过程指标", "T、S、候选质量、token、时间和轮次可追溯", "虚拟指标在留出蓝图/运行中预测真人质量，并报告不确定性"),
            ("真人验证", "至少比较 A/B/C 的信度、汇聚效度、区分效度", "独立确认样本；结构效度；盲态分析；预注册和功效分析"),
            ("内容质量", "专家审核重复、构念纯度和覆盖", "盲审、双人或多人一致性，量化内容质量与心理测量结果的关系"),
            ("效率结论", "报告 token、时间、人工干预和单位合格题成本", "增加 D 人工同任务对照，绘制质量—成本 Pareto 前沿"),
            ("可复现性", "保留原始题库、版本、日志、随机化、分析脚本", "开放代码/匿名数据/材料，或提供受限访问与完整复现实验包"),
        ],
        [1600, 3850, 3910],
        font_size=8.55,
    )

    add_heading(doc, "八、逐日执行计划（2026-09-02—2026-09-30）", 1)
    add_body(doc, "计划按“9 月 15 日完成实验包、9 月 30 日完成初稿”倒排。若伦理审批、招募或模型调用出现硬阻塞，应保留日志并按降级方案调整，不要用未经验证的替代结果填充真人结论。", after=7)
    add_heading(doc, "8.1 实验执行阶段：9 月 2 日—9 月 15 日", 2)
    phase1 = [
        ("9/2", "方案锁定", "锁定 RQ1—RQ4、A/B/C 条件；决定是否加入 D 人工同任务对照；确认最终题数、蓝图和构念。", "一页版研究设计；责任人和截止点"),
        ("9/3", "预注册与测量规则", "完成功效分析；固定 NEO-FFI 与 Mussel 计分；明确 Mussel 0/1 规则、缺失、剔除、随机化和主要比较。", "预注册初稿；数据字典；伦理状态清单"),
        ("9/4", "系统冻结", "冻结模型、prompt、temperature、seed、蓝图输入、题数和输出格式；确保 A/B/C 可独立运行并导出完整日志。", "版本标签；可运行实验包"),
        ("9/5", "冒烟测试", "每个条件完成小规模运行；检查题目数量、重复约束、虚拟施测、局部复测、统一施测、组卷和平台期停止。", "冒烟测试报告；问题清单"),
        ("9/6", "实验一批次 1", "运行 A/B/C 的第一批独立重复；同步启动盲态内容审核。", "原始输出、token/时间日志、审题表"),
        ("9/7", "实验一批次 2", "完成剩余重复与其他蓝图；如有 D，完成专家同任务开发并记录工时。", "实验一完整结果表；运行失败清单"),
        ("9/8", "真人材料冻结", "冻结所有问卷版本；编程施测；做 20—30 人仅用于流程和可理解性检查，不根据结果改主要题目。", "正式问卷链接/程序；试测问题记录"),
        ("9/9", "真人施测启动", "按预注册随机化正式招募；记录同意、完成、缺失、注意力和设备信息。", "招募启动记录；每日数据备份"),
        ("9/10", "质量监控", "检查样本数、完成率和明显技术问题；不查看主要结果差异，不按结果删题或换题。", "数据质量日报"),
        ("9/11", "继续招募", "补足各条件配额；核对 NEO-FFI/Mussel 计分与 SJT 选项编码。", "配额进度；编码核验"),
        ("9/12", "数据收尾准备", "根据预注册目标决定是否关闭招募；冻结主要分析脚本和图表模板。", "盲态分析脚本；数据锁定方案"),
        ("9/13", "数据锁定与主分析", "锁定真人数据；运行信度、汇聚效度、区分效度和预设组间比较。", "主分析表；分析日志"),
        ("9/14", "稳健性检查", "做蓝图/条件分层、留出或确认样本分析；检查虚拟指标与真人指标关系；整理图表。", "稳健性结果；图表 v1"),
        ("9/15", "实验包冻结", "归档代码、prompt、题库版本、原始日志、匿名数据、分析脚本和决策记录；写一页结果摘要。", "实验一/二冻结包；结果摘要"),
    ]
    add_table(doc, ["日期", "主责", "具体工作", "当日交付/门槛"], phase1, [900, 1350, 4700, 2410], font_size=8.05)

    add_heading(doc, "8.2 初稿写作阶段：9 月 16 日—9 月 30 日", 2)
    phase2 = [
        ("9/16", "论文骨架", "建立标题、摘要、引言、方法、结果、讨论、补充材料结构；把每个 RQ 对应到图表。", "论文大纲；图表清单"),
        ("9/17", "研究动机", "写自动测验开发、虚拟被试和真人验证的研究缺口；明确核心贡献边界。", "引言 v1"),
        ("9/18", "相关工作", "整理 LLM 测验生成、自动题目开发、心理测量验证和效率评估文献。", "文献矩阵；引用清单"),
        ("9/19", "理论模型", "写虚拟目标传导度、构念特异度与真人心理测量结果之间的可检验关系。", "理论/假设段落"),
        ("9/20", "系统方法", "描述蓝图、证据、出题、审题、施测、单题局部返修、统一施测、组卷与平台期停止。", "方法第 1 版"),
        ("9/21", "实验一方法", "写消融条件、重复运行、控制变量、输出指标、成本统计和人工审核。", "实验一方法"),
        ("9/22", "实验二方法", "写被试、伦理、随机化、问卷、NEO-FFI、Mussel 0/1 计分和分析方案。", "实验二方法"),
        ("9/23", "实验一结果", "写 A/B/C（和 D）的质量、稳定性、效率、失败率与迭代轨迹；不把过程指标写成真人效度。", "结果第 1 版；图 1—2"),
        ("9/24", "实验二结果", "写真人信度、汇聚效度、区分效度、结构效度和条件比较。", "结果第 2 版；表 1—3"),
        ("9/25", "迁移与成本", "分析虚拟指标—真人指标关系；报告质量—token/时间/人工成本关系与平台期。", "迁移/成本结果；图 3—4"),
        ("9/26", "讨论机制", "解释闭环为何可能有效、哪些机制产生增量、单题局部复测与统一整卷组卷的分工。", "讨论前半"),
        ("9/27", "讨论边界", "写样本、构念数量、模型依赖、社会赞许、虚拟—真人差异、Mussel 对照边界、伦理和可推广性。", "讨论后半；限制清单"),
        ("9/28", "全文成稿", "完成摘要、结论、图表标题、补充材料、方法细节和投稿信要点。", "完整初稿 v1"),
        ("9/29", "内部审阅", "逐项检查统计、分母、效应方向、引用、AI 使用披露、代码/数据可用性和期刊格式。", "修改清单；统计复核"),
        ("9/30", "初稿冻结", "整合意见，冻结可审阅初稿；列出投 BRM 与冲刺 Psychological Methods 所需的补充证据。", "初稿 v2；投稿路线图"),
    ]
    add_table(doc, ["日期", "主责", "具体工作", "当日交付/门槛"], phase2, [900, 1350, 4700, 2410], font_size=8.05)

    add_heading(doc, "九、风险与降级方案", 1)
    add_table(
        doc,
        ["风险", "可能影响", "处理方式"],
        [
            ("伦理审批/监护人同意未完成", "真人实验不能合法启动，9/15 结论延期", "立即确认审批状态；未通过前只完成系统实验和施测程序，不能把模拟结果写成真人证据。"),
            ("招募不足或条件不平衡", "效应估计不稳定，组间比较失真", "预注册配额、延长招募或缩小结论；报告实际效能和缺失，不事后删组。"),
            ("模型接口/调用超时", "实验条件不可复现或成本失控", "固定重试上限、缓存可复用响应、保存中间状态；在论文中报告失败率与实际成本。"),
            ("虚拟指标提高但真人结果不提高", "说明内部指标迁移不足", "把结果作为重要发现；检查指标定义、过拟合、构念污染和样本差异，不强行解释为系统有效。"),
            ("只有一个构念或一个蓝图", "外推性不足，Top 一区风险高", "增加蓝图/构念或明确定位为方法可行性研究，优先投方法匹配度高的期刊。"),
            ("没有人工同任务对照", "无法支持“比人更高效”", "删除该表述；改写为 AI 流程消融 + Mussel 外部参照，并把人工对照列为后续实验。"),
        ],
        [2450, 2700, 4210],
        font_size=8.7,
    )

    add_heading(doc, "十、最终提交包", 1)
    for text in [
        "预注册方案：研究问题、假设、条件、样本量、计分、剔除、主要分析与停止规则。",
        "系统材料：蓝图、evidence、所有 prompt、模型参数、代码版本和运行环境。",
        "实验一数据：每次运行的题库、虚拟作答、单题指标、整卷指标、token、时间、费用、失败/重试和人工干预日志。",
        "实验二数据：匿名真人数据、随机化记录、NEO-FFI 与 Mussel 计分脚本、SJT 计分脚本、缺失和质量控制记录。",
        "分析包：一键或分步可复现的统计脚本、结果表、图表源文件、稳健性分析和版本说明。",
        "论文材料：初稿、补充材料、局限性声明、AI 使用披露、数据/代码可用性声明和目标期刊格式清单。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "参考来源与说明", 1)
    add_body(doc, "期刊定位依据各期刊官方页面；影响因子、JCR 分区和中科院分区会更新，以下链接用于投稿前复核。访问日期：2026-09-03。", size=9.5, color=MUTED)
    add_source(doc, "Behavior Research Methods", "https://www.psychonomic.org/page/BRM", "官方期刊页面")
    add_source(doc, "Psychological Methods", "https://www.apa.org/pubs/journals/met/", "APA 官方期刊页面")
    add_source(doc, "Applied Psychological Measurement", "https://journals.sagepub.com/author-instructions/apm", "SAGE 官方投稿范围页面")
    add_source(doc, "Educational and Psychological Measurement", "https://journals.sagepub.com/home/epm", "SAGE 官方期刊页面")
    add_source(doc, "Psychological Assessment", "https://www.apa.org/pubs/journals/pas", "APA 官方期刊页面")
    add_body(doc, "本文件中的研究设计、阈值和时间安排是当前项目的执行建议；除已锁定的计分规则和实验条件外，其余内容应在预注册与导师确认后冻结。", size=9.2, color=MUTED, italic=True)

    doc.core_properties.title = "LLM 驱动的 SJT 自动开发系统：实验方案、目标期刊与执行计划"
    doc.core_properties.subject = "Top 一区目标下的系统消融实验、真人心理测量验证与 2026 年 9 月执行计划"
    doc.core_properties.author = "研究项目组"
    doc.core_properties.comments = "Generated as a research planning document; verify journal rankings before submission."
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
